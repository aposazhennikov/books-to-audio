"""Loader for DOCX (.docx) book files using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from book_normalizer.loaders.base import BaseLoader
from book_normalizer.models.book import Book, Chapter, Metadata, Paragraph

logger = logging.getLogger(__name__)

_HEADING_STYLES = frozenset({
    "Heading 1", "Heading 2", "Heading 3",
    "Title", "Subtitle",
})


class DocxLoader(BaseLoader):
    """
    DOCX loader built on python-docx.

    Extracts text from document paragraphs and uses Word heading
    styles to identify chapter boundaries. Non-heading paragraphs
    are collected as body text.
    """

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    def load(self, path: Path) -> Book:
        """Load a DOCX file and return a Book."""
        resolved = path.resolve()
        if not resolved.is_file():
            raise FileNotFoundError(f"File not found: {resolved}")

        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX loading. Install it: pip install python-docx"
            ) from exc

        doc = Document(str(resolved))
        metadata = self._extract_metadata(doc, resolved)
        chapters = self._extract_chapters(doc)

        if not chapters:
            paragraphs = self._extract_all_paragraphs(doc)
            chapters = [Chapter(title="Full Text", index=0, paragraphs=paragraphs)]

        book = Book(metadata=metadata, chapters=chapters)
        total_paras = sum(len(ch.paragraphs) for ch in chapters)
        book.add_audit(
            "loading", "docx_loader",
            f"chapters={len(chapters)}, paragraphs={total_paras}",
        )
        return book

    @staticmethod
    def _extract_metadata(doc: object, path: Path) -> Metadata:
        """Extract metadata from DOCX core properties."""
        from docx import Document

        props = doc.core_properties

        title = ""
        author = ""

        if props.title:
            title = props.title
        if props.author:
            author = props.author

        return Metadata(
            title=title or "Untitled",
            author=author or "Unknown",
            source_path=str(path),
            source_format="docx",
        )

    @classmethod
    def _extract_chapters(cls, doc: object) -> list[Chapter]:
        """Split document into chapters based on heading styles."""
        from docx import Document

        chapters: list[Chapter] = []
        current_title = ""
        current_paragraphs: list[Paragraph] = []
        para_idx = 0

        for doc_para in doc.paragraphs:
            style_name = doc_para.style.name if doc_para.style else ""
            text = doc_para.text.strip()

            if not text:
                continue

            is_heading = style_name in _HEADING_STYLES

            if is_heading:
                if current_paragraphs:
                    chapters.append(
                        Chapter(
                            title=current_title or f"Section {len(chapters) + 1}",
                            index=len(chapters),
                            paragraphs=current_paragraphs,
                        )
                    )
                current_title = text
                current_paragraphs = []
                para_idx = 0
            else:
                current_paragraphs.append(
                    Paragraph(
                        raw_text=text,
                        normalized_text="",
                        index_in_chapter=para_idx,
                    )
                )
                para_idx += 1

        if current_paragraphs:
            chapters.append(
                Chapter(
                    title=current_title or f"Section {len(chapters) + 1}",
                    index=len(chapters),
                    paragraphs=current_paragraphs,
                )
            )

        return chapters

    @staticmethod
    def _extract_all_paragraphs(doc: object) -> list[Paragraph]:
        """Fallback: extract all paragraphs as a flat list."""
        from docx import Document

        paragraphs: list[Paragraph] = []
        for idx, doc_para in enumerate(doc.paragraphs):
            text = doc_para.text.strip()
            if text:
                paragraphs.append(
                    Paragraph(raw_text=text, normalized_text="", index_in_chapter=idx)
                )
        return paragraphs
