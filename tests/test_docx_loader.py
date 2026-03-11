"""Tests for the DOCX loader."""

from __future__ import annotations

from pathlib import Path

import pytest

from book_normalizer.loaders.docx_loader import DocxLoader


def _create_test_docx(path: Path, with_headings: bool = True) -> Path:
    """Create a test DOCX file with optional heading-based chapters."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"

    if with_headings:
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("First paragraph of chapter one.")
        doc.add_paragraph("Second paragraph of chapter one.")
        doc.add_heading("Chapter 2", level=1)
        doc.add_paragraph("First paragraph of chapter two.")
    else:
        doc.add_paragraph("Just a flat paragraph.")
        doc.add_paragraph("Another flat paragraph.")

    doc.save(str(path))
    return path


class TestDocxLoader:
    def test_supported_extensions(self) -> None:
        loader = DocxLoader()
        assert ".docx" in loader.supported_extensions

    def test_can_load(self, tmp_path: Path) -> None:
        loader = DocxLoader()
        assert loader.can_load(tmp_path / "doc.docx")
        assert not loader.can_load(tmp_path / "doc.pdf")

    def test_load_nonexistent_file(self, tmp_path: Path) -> None:
        loader = DocxLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(tmp_path / "missing.docx")

    def test_load_with_headings(self, tmp_path: Path) -> None:
        docx_file = _create_test_docx(tmp_path / "with_headings.docx", with_headings=True)

        loader = DocxLoader()
        book = loader.load(docx_file)

        assert book.metadata.title == "Test Document"
        assert book.metadata.author == "Test Author"
        assert book.metadata.source_format == "docx"
        assert len(book.chapters) == 2
        assert book.chapters[0].title == "Chapter 1"
        assert len(book.chapters[0].paragraphs) == 2
        assert book.chapters[1].title == "Chapter 2"
        assert len(book.chapters[1].paragraphs) == 1

    def test_load_without_headings(self, tmp_path: Path) -> None:
        docx_file = _create_test_docx(tmp_path / "no_headings.docx", with_headings=False)

        loader = DocxLoader()
        book = loader.load(docx_file)

        assert len(book.chapters) == 1
        assert len(book.chapters[0].paragraphs) == 2

    def test_paragraph_content(self, tmp_path: Path) -> None:
        docx_file = _create_test_docx(tmp_path / "content.docx", with_headings=True)

        loader = DocxLoader()
        book = loader.load(docx_file)

        assert book.chapters[0].paragraphs[0].raw_text == "First paragraph of chapter one."

    def test_audit_trail(self, tmp_path: Path) -> None:
        docx_file = _create_test_docx(tmp_path / "audit.docx", with_headings=True)

        loader = DocxLoader()
        book = loader.load(docx_file)
        assert any(r["stage"] == "loading" for r in book.audit_trail)
        assert any("docx_loader" in r["action"] for r in book.audit_trail)
